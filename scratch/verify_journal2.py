import docx

doc = docx.Document(r"c:\GWO\Journal2.docx")

keywords_to_check = ["RNN", "Recurrent Neural Network", "40.09", "41.35", "44,46", "46,41"]

with open(r"c:\GWO\scratch\verify_journal2.txt", "w", encoding="utf-8") as f:
    f.write("=== CHECKING FOR LEFTOVER OLD TERMS/NUMBERS ===\n")
    found_issues = False
    for i, p in enumerate(doc.paragraphs):
        for kw in keywords_to_check:
            if kw in p.text:
                f.write(f"[ISSUE] Found '{kw}' in P{i}: {p.text[:100]}\n")
                found_issues = True

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for kw in keywords_to_check:
                    if kw in cell.text:
                        f.write(f"[ISSUE] Found '{kw}' in Table {t_idx+1} Row {r_idx}: {cell.text}\n")
                        found_issues = True

    if not found_issues:
        f.write("SUCCESS: No leftover old terms or numbers found!\n\n")

    f.write("=== PARAGRAPH SUMMARY IN UPDATED JOURNAL2 ===\n")
    for i in [5, 8, 11, 15, 17, 22, 24, 25, 31, 34, 40, 45, 49, 50, 56, 57]:
        f.write(f"\nP{i}: {doc.paragraphs[i].text}\n")

    f.write("\n=== TABLE SUMMARY IN UPDATED JOURNAL2 ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\nTable {t_idx+1}:\n")
        for r in table.rows:
            row_txt = [c.text.strip().replace('\n', ' ') for c in r.cells]
            f.write(" | ".join(row_txt) + "\n")
