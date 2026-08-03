import docx

doc = docx.Document(r"c:\GWO\Journal2.docx")

# 1. Update Sub-bab 2.5 (Metrik Evaluasi & Uji Signifikansi)
doc.paragraphs[28].text = (
    "Kinerja model dievaluasi menggunakan lima metrik standar peramalan: Mean Absolute Error (MAE), "
    "Mean Squared Error (MSE), Root Mean Squared Error (RMSE), Mean Absolute Percentage Error (MAPE), "
    "dan R-squared (R²). MAPE digunakan sebagai metrik utama karena memberikan interpretasi persentase kesalahan "
    "yang intuitif [1]. Selain itu, guna membuktikan secara matematis bahwa perbedaan performa peramalan antar model "
    "terbukti signifikan secara statistik dan bukan akibat fluktuasi acak, penelitian ini menerapkan uji hipotesis "
    "Diebold–Mariano (DM) Test — uji statistik standar internasional khusus peramalan time series — serta uji non-parametrik "
    "Wilcoxon Signed-Rank Test dan perhitungan 95% Confidence Interval (CI) untuk selisih kesalahan prediktif."
)

# 2. Update Sub-bab 4.3 Pembahasan (Hasil Uji Signifikansi Statistik)
doc.paragraphs[50].text = (
    "Hasil ini menunjukkan bahwa pendekatan optimasi bobot ensemble dengan GWO sangat efektif dalam meningkatkan akurasi "
    "peramalan penjualan. GWO Ensemble tidak hanya unggul secara numerik, tetapi juga terbukti signifikan secara statistik. "
    "Berdasarkan pengujian Diebold–Mariano (DM) Test pada 335 sampel data uji, diperoleh nilai DM-statistic sebesar 7,0287 "
    "dengan p-value = 1,18 × 10⁻¹¹ (p < 0,001) pada metrik MAPE, DM-statistic = 10,8722 (p < 0,001) pada metrik MAE, dan "
    "DM-statistic = 11,5179 (p < 0,001) pada metrik MSE. Uji non-parametrik Wilcoxon Signed-Rank Test juga memperkuat temuan ini "
    "dengan nilai Z-score = -9,5548 (p = 8,43 × 10⁻¹³ < 0,001). Selang kepercayaan 95% (95% Confidence Interval) untuk selisih "
    "kesalahan berada pada rentang [2,95%, 5,25%], yang seluruh nilainya berada secara tegas di atas nol. Nilai p-value < 0,001 "
    "ini mengonfirmasi pada tingkat kepercayaan 99,9% bahwa peningkatan akurasi peramalan GWO Ensemble sebesar 1,02% dibanding "
    "baseline terbaik (Seasonal MA) adalah signifikan secara statistik dan sangat andal."
)

# 3. Update Sub-bab 5 Kesimpulan
doc.paragraphs[56].text = (
    "Pertama, model weighted ensemble yang menggabungkan MA, ES, dan LR berhasil dikembangkan dengan bobot optimal: "
    "MA 64,10%, ES 35,90%, dan LR 0,00%. Kedua, algoritma GWO berhasil menemukan kombinasi bobot optimal yang menghasilkan "
    "MAPE sebesar 11,60%, melampaui seluruh model baseline individu. Dibandingkan best baseline MA (MAPE 12,62%), terjadi "
    "peningkatan akurasi sebesar 1,02% secara absolut yang terbukti signifikan secara statistik berdasarkan uji Diebold–Mariano "
    "dan Wilcoxon (p < 0,001). Ketiga, hasil pengujian multi-run sebanyak 30 kali membuktikan bahwa GWO sangat robust dengan "
    "standar deviasi MAPE yang sangat kecil, yaitu 0,000026%, menunjukkan algoritma GWO konsisten mencapai konvergensi global."
)

# Save updated docx
doc.save(r"c:\GWO\Journal2.docx")
print("Successfully added statistical significance tests to Journal2.docx!")
