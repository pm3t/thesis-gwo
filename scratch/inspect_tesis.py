import docx

doc = docx.Document(r"c:\GWO\TesisV2_backup.docx")

with open(r"c:\GWO\scratch\tesis_dump.txt", "w", encoding="utf-8") as f:
    f.write(f"Total Paragraphs: {len(doc.paragraphs)}\n")
    f.write(f"Total Tables: {len(doc.tables)}\n\n")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") or "BAB" in p.text or "Tabel" in p.text or "Gambar" in p.text:
            f.write(f"P{i} [{p.style.name}]: {p.text}\n")
