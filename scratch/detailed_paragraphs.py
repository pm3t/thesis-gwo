import docx

doc = docx.Document(r"c:\GWO\Journal2.docx")

with open(r"c:\GWO\scratch\detailed_paragraphs.txt", "w", encoding="utf-8") as f:
    for i, p in enumerate(doc.paragraphs):
        f.write(f"--- P{i} [{p.style.name}] ---\n{p.text}\n\n")

    f.write("\n=== TABLES ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {t_idx+1} ---\n")
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            f.write(f"Row {r_idx}: {' | '.join(cells)}\n")
