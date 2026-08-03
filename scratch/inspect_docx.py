import docx

doc = docx.Document(r"c:\GWO\Journal2.docx")

with open(r"c:\GWO\scratch\journal2_dump.txt", "w", encoding="utf-8") as f:
    f.write("=== PARAGRAPHS IN JOURNAL2.DOCX ===\n")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            f.write(f"P{i} [{p.style.name}]: {p.text}\n\n")

    f.write("\n=== TABLES IN JOURNAL2.DOCX ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\nTable {t_idx+1}:\n")
        for r in table.rows:
            row_txt = [cell.text.strip().replace('\n', ' ') for cell in r.cells]
            f.write(" | ".join(row_txt) + "\n")
