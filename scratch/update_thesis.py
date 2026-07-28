import docx
import pandas as pd
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document('TesisV2.docx')

# 1. Update Table 24 (Table 4.8 Baseline Performance)
t24 = doc.tables[24]
t24.rows[0].cells[1].text = 'Moving Average (MA)'
t24.rows[0].cells[2].text = 'Exponential Smoothing (ES)'
t24.rows[0].cells[3].text = 'Linear Regression (LR)'

t24.rows[1].cells[1].text = '12.6202%'
t24.rows[1].cells[2].text = '13.8889%'
t24.rows[1].cells[3].text = '20.2528%'

t24.rows[2].cells[1].text = '0.0722'
t24.rows[2].cells[2].text = '0.0793'
t24.rows[2].cells[3].text = '0.0987'

t24.rows[3].cells[1].text = '0.0102'
t24.rows[3].cells[2].text = '0.0116'
t24.rows[3].cells[3].text = '0.0132'

t24.rows[4].cells[1].text = '0.1012'
t24.rows[4].cells[2].text = '0.1079'
t24.rows[4].cells[3].text = '0.1148'

t24.rows[5].cells[1].text = '0.6149'
t24.rows[5].cells[2].text = '0.5623'
t24.rows[5].cells[3].text = '0.5046'

# 2. Update Table 25 (Table 4.9 30-run Optimasi GWO)
t25 = doc.tables[25]
df_runs = pd.read_csv('Analysis/Stability_Runs_GWO.csv')

# Header
t25.rows[0].cells[0].text = 'Run'
t25.rows[0].cells[1].text = 'n (wolf)'
t25.rows[0].cells[2].text = 'Iteration'
t25.rows[0].cells[3].text = 'MAPE (%)'
t25.rows[0].cells[4].text = 'w1 (MA)'
t25.rows[0].cells[5].text = 'w2 (ES)'
t25.rows[0].cells[6].text = 'w3 (LR)'

for i, row_data in df_runs.iterrows():
    if i + 1 < len(t25.rows):
        r = t25.rows[i + 1]
        r.cells[0].text = str(int(row_data['Run']))
        r.cells[1].text = str(int(row_data['n (wolf)']))
        r.cells[2].text = str(int(row_data['iteration']))
        r.cells[3].text = f"{row_data['MAPE']:.4f}%"
        r.cells[4].text = f"{row_data['w1']:.4f}"
        r.cells[5].text = f"{row_data['w2']:.4f}"
        r.cells[6].text = f"{row_data['w3']:.4f}"

# 3. Update Table 26 (Table 4.10 Stability Summary)
t26 = doc.tables[26]
t26.rows[0].cells[0].text = 'Parameter Statistik'
t26.rows[0].cells[1].text = 'Nilai Fitness (MAPE)'
t26.rows[0].cells[2].text = 'Bobot w1 (MA)'
t26.rows[0].cells[3].text = 'Bobot w2 (ES)'
t26.rows[0].cells[4].text = 'Bobot w3 (LR)'

t26.rows[1].cells[0].text = 'Terbaik'
t26.rows[1].cells[1].text = '11.6045%'
t26.rows[1].cells[2].text = '0.6410'
t26.rows[1].cells[3].text = '0.3590'
t26.rows[1].cells[4].text = '0.0000'

t26.rows[2].cells[0].text = 'Terburuk'
t26.rows[2].cells[1].text = '11.6047%'
t26.rows[2].cells[2].text = '0.6409'
t26.rows[2].cells[3].text = '0.3589'
t26.rows[2].cells[4].text = '0.0002'

t26.rows[3].cells[0].text = 'Rata-rata'
t26.rows[3].cells[1].text = '11.6045%'
t26.rows[3].cells[2].text = '0.6410'
t26.rows[3].cells[3].text = '0.3590'
t26.rows[3].cells[4].text = '0.0000'

t26.rows[4].cells[0].text = 'Std. Deviation'
t26.rows[4].cells[1].text = '0.000052%'
t26.rows[4].cells[2].text = '0.000038'
t26.rows[4].cells[3].text = '0.000033'
t26.rows[4].cells[4].text = '0.000039'

# 4. Update Table 27 (Table 4.11 Comparison Overview)
t27 = doc.tables[27]
t27.rows[0].cells[0].text = 'Model'
t27.rows[0].cells[1].text = 'MAPE (%)'
t27.rows[0].cells[2].text = 'MAE'
t27.rows[0].cells[3].text = 'MSE'
t27.rows[0].cells[4].text = 'RMSE'
t27.rows[0].cells[5].text = 'R²'

# Row 1: MA
t27.rows[1].cells[0].text = 'MA'
t27.rows[1].cells[1].text = '12.6202%'
t27.rows[1].cells[2].text = '0.0722'
t27.rows[1].cells[3].text = '0.0102'
t27.rows[1].cells[4].text = '0.1012'
t27.rows[1].cells[5].text = '0.6149'

# Row 2: ES
t27.rows[2].cells[0].text = 'ES'
t27.rows[2].cells[1].text = '13.8889%'
t27.rows[2].cells[2].text = '0.0793'
t27.rows[2].cells[3].text = '0.0116'
t27.rows[2].cells[4].text = '0.1079'
t27.rows[2].cells[5].text = '0.5623'

# Row 3: LR
t27.rows[3].cells[0].text = 'LR'
t27.rows[3].cells[1].text = '20.2528%'
t27.rows[3].cells[2].text = '0.0987'
t27.rows[3].cells[3].text = '0.0132'
t27.rows[3].cells[4].text = '0.1148'
t27.rows[3].cells[5].text = '0.5046'

# Row 4: GWO Ensemble
if len(t27.rows) > 4:
    t27.rows[4].cells[0].text = 'GWO Ensemble'
    t27.rows[4].cells[1].text = '11.6045%'
    t27.rows[4].cells[2].text = '0.0663'
    t27.rows[4].cells[3].text = '0.0083'
    t27.rows[4].cells[4].text = '0.0913'
    t27.rows[4].cells[5].text = '0.6865'

# 5. Update Paragraphs
doc.paragraphs[531].text = 'Gambar 4.3 menyajikan perbandingan antara hasil prediksi model Seasonal Moving Average (MA) dengan data aktual pada periode uji. Garis hitam menunjukkan data aktual, sedangkan garis biru menunjukkan hasil prediksi MA. Model MA berhasil menangkap komponen tren dan musiman tingkat tinggi dengan sangat baik, menghasilkan tingkat kesalahan MAPE sebesar 12,6202% yang menjadikannya model baseline individu terbaik.'
doc.paragraphs[535].text = 'Gambar 4.4 menunjukkan perbandingan prediksi model Holt-Winters Exponential Smoothing (ES) terhadap data aktual. Model ES mampu mengikuti pergerakan tren dan musiman data dengan nilai MAPE sebesar 13,8889%, MAE 0,0793, RMSE 0,1079, serta koefisien determinasi R² sebesar 0,5623.'
doc.paragraphs[538].text = 'Gambar 4.5 memperlihatkan perbandingan prediksi model Linear Regression (LR) terhadap data aktual. Model Linear Regression (LR) memodelkan hubungan tren linier dan prediktor waktu harian/bulanan, menghasilkan nilai MAPE sebesar 20,2528%, MAE 0,0987, RMSE 0,1148, dan R² sebesar 0,5046.'

doc.paragraphs[540].text = 'Berdasarkan data pada Tabel 4.8, model Seasonal Moving Average (MA) mencatatkan performa terbaik di antara ketiga model baseline individu dengan nilai MAPE terendah, yakni 12,6202%, diikuti oleh Holt-Winters Exponential Smoothing (ES) sebesar 13,8889%, dan Linear Regression (LR) sebesar 20,2528%.'
doc.paragraphs[541].text = 'Model MA juga menghasilkan nilai koefisien determinasi R² tertinggi (0,6149) serta nilai kesalahan MAE (0,0722) dan RMSE (0,1012) yang paling kecil di antara seluruh model baseline individu. Hal ini menunjukkan bahwa pola musiman mingguan pada dataset penjualan toko dapat ditangkap secara sangat baik oleh metode rata-rata bergerak musiman.'

doc.paragraphs[550].text = 'Berdasarkan ringkasan statistik hasil optimasi GWO di Tabel 4.10, Algoritma GWO mencapai nilai fitness (MAPE) terbaik sebesar 11,604454%, MAPE terburuk sebesar 11,604725%, dan rata-rata dari nilai MAPE dari 30 run sebesar 11,604476%. Standar deviasi MAPE berada pada angka yang sangat kecil, yaitu 0,000052%. Nilai deviasi yang mendekati nol ini membuktikan stabilitas dan konsistensi algoritma GWO yang sangat tinggi.'
doc.paragraphs[551].text = 'Pada run dengan performa terbaik, konfigurasi bobot optimal yang dihasilkan adalah w1 (MA) = 0,640968, w2 (ES) = 0,359032, dan w3 (LR) = 0,000000. Konfigurasi ini konsisten dengan rata-rata bobot 30 run, di mana rata-rata bobot w1 (MA) = 0,640956, w2 (ES) = 0,359029, dan w3 (LR) = 0,000014.'

doc.paragraphs[557].text = 'Kombinasi bobot optimal memperlihatkan dominasi bobot Seasonal Moving Average (MA) sebesar 64,10% (0,640968), diikuti oleh Holt-Winters Exponential Smoothing (ES) sebesar 35,90% (0,359032), sementara Linear Regression (LR) mendapatkan bobot 0,00% (0,000000). Hal ini mengindikasikan bahwa GWO secara efisien memprioritaskan dua model baseline berkinerja terbaik (MA dan ES) untuk digabungkan, serta mengeliminasi kontribusi model dengan kesalahan lebih tinggi (LR) guna meminimalkan fungsi objektif MAPE.'
doc.paragraphs[559].text = 'Model Seasonal Moving Average (MA) memberikan kontribusi paling dominan sebesar 64,10% dari total bobot ensemble, disusul oleh Holt-Winters Exponential Smoothing (ES) sebesar 35,90%, sedangkan Linear Regression (LR) tidak digunakan (bobot 0,00%).'

doc.paragraphs[564].text = 'Berdasarkan komparasi data pada Tabel 4.11, model GWO Ensemble menghasilkan kinerja terbaik di antara seluruh model dengan MAPE sebesar 11,6045%, MAE 0,0663, MSE 0,0083, RMSE 0,0913, dan R² 0,6865. Apabila disandingkan dengan performa model individu terbaik (best baseline) yaitu MA (MAPE 12,6202%), model GWO Ensemble berhasil menurunkan kesalahan prediksi sebesar 1,0157% secara absolut, yang setara dengan peningkatan akurasi (accuracy improvement) sebesar 8,05%.'
doc.paragraphs[567].text = 'Gambar 4.7 menyajikan perbandingan nilai MAPE dari seluruh model yang diuji dalam penelitian ini, yaitu Seasonal Moving Average (MA: 12,62%), Holt-Winters Exponential Smoothing (ES: 13,89%), Linear Regression (LR: 20,25%), dan GWO Ensemble (11,60%). Grafik ini menegaskan bahwa GWO Ensemble mencapai MAPE paling rendah di antara semua model.'
doc.paragraphs[570].text = 'Gambar 4.8 memperlihatkan perbandingan antara hasil prediksi model GWO Ensemble dengan data aktual pada periode pengujian. Dengan MAPE 11,60%, GWO Ensemble mampu melacak pola pergerakan penjualan aktual dengan presisi yang lebih tinggi dibandingkan seluruh model individu.'
doc.paragraphs[571].text = 'Peningkatan Akurasi Model GWO Ensemble'
doc.paragraphs[572].text = 'Peningkatan akurasi sebesar 8,05% yang dicapai oleh GWO Ensemble membuktikan bahwa algoritma GWO berhasil mengombinasikan keunggulan MA (64,10%) dan ES (35,90%) untuk menghasilkan prediksi ensemble yang lebih akurat dibandingkan model baseline tunggal mana pun.'

# Bab V Kesimpulan
doc.paragraphs[580].text = 'Penelitian ini telah berhasil mengembangkan model weighted ensemble yang menggabungkan tiga model baseline — Seasonal Moving Average (MA), Holt-Winters Exponential Smoothing (ES), dan Linear Regression (LR). Hasil optimasi GWO menghasilkan kombinasi bobot optimal w1 (MA) = 0,640968 (64,10%), w2 (ES) = 0,359032 (35,90%), dan w3 (LR) = 0,000000 (0,00%). Model MA memberikan kontribusi terbesar karena kinerjanya yang paling baik di antara model baseline tunggal, dilengkapi oleh ES untuk memberikan stabilitas pemulusan.'
doc.paragraphs[581].text = 'Algoritma Grey Wolf Optimizer (GWO) berhasil menemukan kombinasi bobot optimal yang meminimalkan kesalahan prediksi, menghasilkan GWO Ensemble dengan MAPE sebesar 11,6045%, MAE 0,0663, RMSE 0,0913, dan R² 0,6865. Dibandingkan model baseline terbaik (MA dengan MAPE 12,6202%), terjadi peningkatan akurasi sebesar 8,05%.'
doc.paragraphs[582].text = 'Model GWO Ensemble terbukti unggul secara menyeluruh di semua metrik evaluasi (MAPE 11,6045% vs MA 12,6202%, MAE 0,0663 vs MA 0,0722, RMSE 0,0913 vs MA 0,1012, dan R² 0,6865 vs MA 0,6149), mengonfirmasi bahwa pembobotan optimal GWO mampu mengungguli seluruh model baseline individu secara konsisten.'
doc.paragraphs[583].text = 'Hasil pengujian 30 run membuktikan stabilitas algoritma GWO yang sangat tinggi dengan standar deviasi MAPE sebesar 0,000052%, menunjukkan bahwa GWO konsisten mencapai solusi optimal global pada setiap percobaan.'

doc.save('TesisV2.docx')
print('Bab IV & V updated successfully with exact analysis data.')
